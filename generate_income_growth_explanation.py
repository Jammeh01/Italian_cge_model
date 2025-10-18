"""
Generate Word Document Explanation for Regional Income Growth Graphs
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_explanation_document():
    """Create a detailed explanation document for the regional income growth visualizations"""

    doc = Document()

    # Title
    title = doc.add_heading(
        'Regional Household Income Growth Analysis: CGE Model Projections (2021-2040)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        'Economic Interpretation and Policy Implications')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Main text content (200 words)
    main_text = """The Computable General Equilibrium (CGE) model projections reveal significant regional disparities in household income growth across Italy under three climate policy scenarios. Under the Business-as-Usual (BAU) scenario, Southern regions exhibit the strongest income growth, with the South achieving 47.5% (€295.2B to €435.6B) and Islands 45.5% growth by 2040, reflecting catch-up dynamics and structural convergence (Dixon & Rimmer, 2002). Northern regions show more modest growth—Northwest 20.7% and Northeast 19.9%—consistent with mature economy characteristics.

The ETS1 scenario (Emissions Trading System Phase 1) demonstrates carbon pricing impacts on regional income trajectories, with slightly reduced growth rates compared to BAU: South 44.9%, Islands 45.6%, suggesting minimal adverse effects on less industrialized regions (Böhringer et al., 2012). The Centre experiences 26.9% growth, indicating moderate policy burden. The ETS2 scenario (2027-2040) shows compressed growth periods but maintains positive trajectories across all regions, with Islands leading at 27.7% growth.

These findings suggest that carbon pricing mechanisms can be implemented without compromising long-term income convergence objectives (Carbone et al., 2009). Policy implications include the need for complementary regional development strategies to maintain growth momentum while achieving environmental targets. The model indicates that poorer regions' growth patterns remain resilient under climate policy frameworks, supporting arguments for integrated environmental-economic policy design (Capros et al., 2013)."""

    # Add main text
    main_para = doc.add_paragraph(main_text)
    main_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in main_para.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # References section
    ref_heading = doc.add_heading('References', 1)

    # List of references
    references = [
        'Böhringer, C., Rutherford, T. F., & Tol, R. S. (2009). THE EU 20/20/2020 targets: An overview of the EMF22 assessment. Energy Economics, 31(Supplement 2), S268-S273.',

        'Capros, P., Paroussos, L., Fragkos, P., Tsani, S., Boitier, B., Wagner, F., ... & Bollen, J. (2013). European decarbonisation pathways under alternative technological and policy choices: A multi-model analysis. Energy Strategy Reviews, 2(3-4), 231-245.',

        'Carbone, J. C., Helm, C., & Rutherford, T. F. (2009). The case for international emission trade in the absence of cooperative climate policy. Journal of Environmental Economics and Management, 58(3), 266-280.',

        'Dixon, P. B., & Rimmer, M. T. (2002). Dynamic General Equilibrium Modelling for Forecasting and Policy: A Practical Guide and Documentation of MONASH. North-Holland: Amsterdam.'
    ]

    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)
        ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in ref_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Save document
    output_path = 'results/Regional_Income_Growth_Explanation.docx'
    doc.save(output_path)

    print("="*80)
    print("✅ Word Document Created Successfully!")
    print("="*80)
    print(f"File saved to: {output_path}")
    print(f"Word count: ~200 words")
    print("\nDocument includes:")
    print("  - Title and subtitle")
    print("  - Detailed economic interpretation")
    print("  - Policy implications")
    print("  - In-text citations")
    print("  - Complete reference list (4 sources)")
    print("="*80)


if __name__ == "__main__":
    create_explanation_document()
