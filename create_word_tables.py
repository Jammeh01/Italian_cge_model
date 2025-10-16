"""
Generate Microsoft Word document with CGE Model Scenario Comparison Tables
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_table_borders(table):
    """Add borders to all cells in a table"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Create table borders element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_scenario_tables_word():
    """Create Word document with all scenario comparison tables"""

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading('Comparative Tables: CGE Model Scenario Setup', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(
        'Italian CGE Model - BAU, ETS1, and ETS2 Scenarios', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========================================================================
    # TABLE 1: Overview of Scenario Definitions
    # ========================================================================
    doc.add_heading('Table 1: Overview of Scenario Definitions', level=2)

    table1 = doc.add_table(rows=4, cols=5)
    table1.style = 'Light Grid Accent 1'

    # Header row
    headers1 = ['Scenario', 'Full Name', 'Time Period',
                'Policy Description', 'Primary Objective']
    for i, header in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data rows
    data1 = [
        ['BAU', 'Business as Usual', '2021-2050', 'No additional climate policies beyond existing regulations',
            'Baseline reference trajectory without carbon pricing'],
        ['ETS1', 'EU ETS Phase 4', '2021-2050', 'EU Emissions Trading System Phase 4 for industrial sectors',
            'Decarbonize industry through carbon pricing with Market Stability Reserve (MSR)'],
        ['ETS2', 'EU ETS Phase 4 + Buildings/Transport',
            '2027-2050 (ETS2 component)', 'Extended ETS coverage including buildings and transport sectors', 'Comprehensive decarbonization across industry, buildings, and transport']
    ]

    for i, row_data in enumerate(data1, start=1):
        for j, value in enumerate(row_data):
            table1.rows[i].cells[j].text = value

    add_table_borders(table1)
    doc.add_paragraph()

    # ========================================================================
    # TABLE 2: Carbon Pricing Parameters
    # ========================================================================
    doc.add_heading('Table 2: Carbon Pricing Parameters', level=2)

    table2 = doc.add_table(rows=9, cols=4)
    table2.style = 'Light Grid Accent 1'

    # Header
    headers2 = ['Parameter', 'BAU', 'ETS1', 'ETS2']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data
    data2 = [
        ['Carbon Price (2021)', '€0/tCO₂e', '€53.90/tCO₂e',
         '€53.90/tCO₂e (ETS1)'],
        ['Carbon Price (2027)', '€0/tCO₂e', '€69.72/tCO₂e',
         '€69.72/tCO₂e (ETS1)\n€45.00/tCO₂e (ETS2)'],
        ['Price Growth Rate', 'N/A', '4.0% annually',
            '4.0% (ETS1)\n2.5% (ETS2)'],
        ['Growth Rate Decline', 'N/A', '0.15% per year',
            '0.15% (ETS1)\n0.10% (ETS2)'],
        ['Maximum Price', 'N/A',
            '€150/tCO₂e (practical limit)', '€150/tCO₂e (ETS1)\n€100/tCO₂e (ETS2)'],
        ['Price Mechanism', 'N/A',
            'Market Stability Reserve (MSR)', 'MSR (ETS1)\nPrice Stability Mechanism (PSM) (ETS2)'],
        ['Price Floor', 'N/A', 'None (MSR managed)',
         'None (ETS1)\n€22.00/tCO₂e (ETS2)'],
        ['Price Cap', 'N/A', 'None (MSR prevents extremes)',
         'None (ETS1)\n€45.00/tCO₂e (ETS2)']
    ]

    for i, row_data in enumerate(data2, start=1):
        for j, value in enumerate(row_data):
            table2.rows[i].cells[j].text = value

    add_table_borders(table2)
    doc.add_paragraph()

    # ========================================================================
    # TABLE 3: Sectoral Coverage
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('Table 3: Sectoral Coverage', level=2)

    table3 = doc.add_table(rows=10, cols=5)
    table3.style = 'Light Grid Accent 1'

    # Header
    headers3 = ['Sector Code', 'Sector Name', 'BAU', 'ETS1', 'ETS2']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data
    sectors_data = [
        ['IND', 'Industry (Manufacturing)', '❌', '✅', '✅'],
        ['GAS', 'Gas Supply', '❌', '✅', '✅'],
        ['OENERGY', 'Other Energy', '❌', '✅', '✅'],
        ['AIR', 'Air Transport', '❌', '✅', '✅'],
        ['WATER', 'Water Transport', '❌', '✅', '✅'],
        ['ROAD', 'Road Transport', '❌', '❌', '✅'],
        ['OTRANS', 'Other Transport', '❌', '❌', '✅'],
        ['SERVICES', 'Services (Buildings)', '❌', '❌', '✅'],
        ['Total Covered', '-', '0', '5 sectors', '8 sectors']
    ]

    for i, row_data in enumerate(sectors_data, start=1):
        for j, value in enumerate(row_data):
            table3.rows[i].cells[j].text = value

    add_table_borders(table3)

    # Add legend
    legend = doc.add_paragraph()
    legend.add_run('Coverage Legend: ').bold = True
    legend.add_run('✅ = Covered by ETS carbon pricing | ❌ = Not covered')

    doc.add_paragraph()

    # ========================================================================
    # TABLE 4: Free Allowance Allocation
    # ========================================================================
    doc.add_heading('Table 4: Free Allowance Allocation', level=2)

    table4 = doc.add_table(rows=7, cols=5)
    table4.style = 'Light Grid Accent 1'

    # Header
    headers4 = ['Parameter', 'BAU', 'ETS1',
                'ETS2 (ETS1 component)', 'ETS2 (ETS2 component)']
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data
    data4 = [
        ['Initial Free Allocation Rate', 'N/A', '80%', '80%', '60%'],
        ['Annual Decline Rate', 'N/A', '2% per year', '2% per year', '3% per year'],
        ['Minimum Allocation', 'N/A', '10%', '10%', '5%'],
        ['Free Allocation (2021)', 'N/A', '80%', '80%', 'N/A (starts 2027)'],
        ['Free Allocation (2030)', 'N/A', '62%', '62%', '51%'],
        ['Free Allocation (2050)', 'N/A', '22% (capped at 10%)',
         '22% (capped at 10%)', '9% (capped at 5%)']
    ]

    for i, row_data in enumerate(data4, start=1):
        for j, value in enumerate(row_data):
            table4.rows[i].cells[j].text = value

    add_table_borders(table4)
    doc.add_paragraph()

    # ========================================================================
    # TABLE 5: Energy Efficiency and Technology Parameters
    # ========================================================================
    doc.add_page_break()
    doc.add_heading(
        'Table 5: Energy Efficiency and Technology Parameters', level=2)

    table5 = doc.add_table(rows=8, cols=4)
    table5.style = 'Light Grid Accent 1'

    # Header
    headers5 = ['Parameter', 'BAU', 'ETS1', 'ETS2']
    for i, header in enumerate(headers5):
        cell = table5.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data
    data5 = [
        ['AEEI (Covered Sectors)', '1.0% per year',
         '1.5% per year', '1.5% per year'],
        ['AEEI (Non-Covered)', '1.0% per year',
         '1.0% per year', '1.0% per year'],
        ['Electrification Rate Growth', '2.5% per year',
            '2.5% per year', '2.5% per year'],
        ['Renewable Share Growth', '4.5% per year',
            '4.5% per year', '4.5% per year'],
        ['Renewable Investment Acceleration',
            'None (baseline)', '20% boost', '40% boost (industry)\n60% boost (South/Islands)'],
        ['Renewable Capacity (2021)', '60 GW', '60 GW', '60 GW'],
        ['Target Renewable Share (2050)', '~70%', '~80%', '~90%']
    ]

    for i, row_data in enumerate(data5, start=1):
        for j, value in enumerate(row_data):
            table5.rows[i].cells[j].text = value

    add_table_borders(table5)

    # Note
    note = doc.add_paragraph()
    note.add_run('Note: ').bold = True
    note.add_run('AEEI = Autonomous Energy Efficiency Improvement')
    doc.add_paragraph()

    # ========================================================================
    # TABLE 6: Regional GDP Growth Rates
    # ========================================================================
    doc.add_heading(
        'Table 6: Regional GDP Growth Rates (Base Assumptions)', level=2)

    table6 = doc.add_table(rows=6, cols=4)
    table6.style = 'Light Grid Accent 1'

    # Header
    headers6 = ['Region', 'BAU Growth Rate',
                'ETS1 Adjustment', 'ETS2 Adjustment']
    for i, header in enumerate(headers6):
        cell = table6.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data
    data6 = [
        ['Northwest', '1.5% per year',
            '-0.4% (industrial costs)', '-0.2% (offset by green investment)'],
        ['Northeast', '1.8% per year',
            '-0.4% (industrial costs)', '-0.2% (offset by green investment)'],
        ['Centre', '1.6% per year',
            '+0.3% (green services)', '+0.4% (green buildings)'],
        ['South', '2.2% per year',
            '+0.3% (green investment)', '+0.4% (renewable projects)'],
        ['Islands', '2.0% per year',
            '+0.3% (green investment)', '+0.4% (renewable projects)']
    ]

    for i, row_data in enumerate(data6, start=1):
        for j, value in enumerate(row_data):
            table6.rows[i].cells[j].text = value

    add_table_borders(table6)
    doc.add_paragraph()

    # ========================================================================
    # TABLE 7: Expected Scenario Outcomes (2050)
    # ========================================================================
    doc.add_page_break()
    doc.add_heading(
        'Table 7: Expected Scenario Outcomes (2050 vs. Baseline)', level=2)

    table7 = doc.add_table(rows=8, cols=4)
    table7.style = 'Light Grid Accent 1'

    # Header
    headers7 = ['Indicator', 'BAU (2050)', 'ETS1 (2050)', 'ETS2 (2050)']
    for i, header in enumerate(headers7):
        cell = table7.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data
    data7 = [
        ['GDP Impact', 'Baseline', '-0.5% to -1.0%', '-1.0% to -1.5%'],
        ['CO₂ Emissions Reduction',
            '0% (baseline)', '-15% to -20%', '-35% to -45%'],
        ['CO₂ Intensity Reduction', 'Natural decline',
            '-25% to -30%', '-45% to -55%'],
        ['Renewable Electricity Share', '~70%', '~80%', '~90%'],
        ['Carbon Revenue (€ billion)', '€0', '€15-20', '€25-35'],
        ['Renewable Capacity (GW)', '~140', '~160', '~200'],
        ['Employment Effect (million)', 'Baseline',
         '-0.1 to 0.0', '0.0 to +0.2']
    ]

    for i, row_data in enumerate(data7, start=1):
        for j, value in enumerate(row_data):
            table7.rows[i].cells[j].text = value

    add_table_borders(table7)
    doc.add_paragraph()

    # ========================================================================
    # TABLE 8: Model Technical Specifications
    # ========================================================================
    doc.add_heading('Table 8: Model Technical Specifications', level=2)

    table8 = doc.add_table(rows=11, cols=2)
    table8.style = 'Light Grid Accent 1'

    # Header
    table8.rows[0].cells[0].text = 'Feature'
    table8.rows[0].cells[1].text = 'Specification'
    for cell in table8.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D9E2F3')

    # Data
    data8 = [
        ['Model Type', 'Recursive dynamic CGE'],
        ['Time Horizon', '2021-2050 (30 years, annual)'],
        ['Sectoral Aggregation',
            '5 aggregate sectors (Agriculture, Industry, Energy, Transport, Services)'],
        ['Regional Disaggregation',
            '5 macro-regions (Northwest, Northeast, Centre, South, Islands)'],
        ['Household Types', '5 regional households'],
        ['Energy Carriers', '3 types (Electricity, Gas, Other Energy)'],
        ['Production Technology',
            'Nested CES functions (VA-Energy-Materials)'],
        ['Trade Specification', 'Armington (imports) + CET (exports)'],
        ['Solution Algorithm', 'IPOPT nonlinear solver'],
        ['Equilibrium Concept', 'Walrasian general equilibrium']
    ]

    for i, row_data in enumerate(data8, start=1):
        table8.rows[i].cells[0].text = row_data[0]
        table8.rows[i].cells[1].text = row_data[1]

    add_table_borders(table8)
    doc.add_paragraph()

    # ========================================================================
    # Summary Section
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('Summary of Key Differences', level=2)

    # BAU
    doc.add_heading('BAU (Business as Usual):', level=3)
    bau_list = doc.add_paragraph(style='List Bullet')
    bau_list.add_run('No carbon pricing')
    doc.add_paragraph('Natural technological progress only',
                      style='List Bullet')
    doc.add_paragraph('Renewable growth from market forces',
                      style='List Bullet')
    doc.add_paragraph('Serves as baseline reference', style='List Bullet')

    # ETS1
    doc.add_heading('ETS1 (EU ETS Phase 4):', level=3)
    doc.add_paragraph('Industrial carbon pricing from 2021',
                      style='List Bullet')
    doc.add_paragraph(
        'Initial price: €53.90/tCO₂, growing at 4% p.a.', style='List Bullet')
    doc.add_paragraph(
        'Market Stability Reserve (no hard price cap)', style='List Bullet')
    doc.add_paragraph('51% emissions coverage', style='List Bullet')
    doc.add_paragraph(
        'Moderate renewable investment boost (20%)', style='List Bullet')

    # ETS2
    doc.add_heading('ETS2 (Comprehensive ETS):', level=3)
    doc.add_paragraph(
        'Extended to buildings and transport from 2027', style='List Bullet')
    doc.add_paragraph(
        'Dual price system: ETS1 + ETS2 (€45/tCO₂ initial)', style='List Bullet')
    doc.add_paragraph(
        'Price Stability Mechanism with €45 ceiling, €22 floor', style='List Bullet')
    doc.add_paragraph('75% emissions coverage', style='List Bullet')
    doc.add_paragraph(
        'Strong renewable investment boost (40-60%)', style='List Bullet')
    doc.add_paragraph(
        'Greater regional equity (benefits South/Islands)', style='List Bullet')

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run('Source: ').italic = True
    footer_para.add_run(
        'Italian CGE Model Configuration Files (2024)').italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_file = 'CGE_Scenario_Comparison_Tables.docx'
    doc.save(output_file)
    print(f"✓ Word document created successfully: {output_file}")
    return output_file


if __name__ == "__main__":
    print("Creating Word document with scenario comparison tables...")
    print("="*60)
    create_scenario_tables_word()
    print("="*60)
    print("Done! The document is ready for use in your paper.")
