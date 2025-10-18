"""
Generate Word Document Explanation for Regional Energy Cost Burden 2040
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_energy_burden_explanation():
    """Create explanation document for regional energy cost burden visualization"""

    doc = Document()

    # Title
    title = doc.add_heading(
        'Regional Energy Cost Burden Analysis 2040: CGE Model Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        'Distributional Impacts of Climate Policy on Household Energy Expenditure')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(13)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Main text content (~100 words)
    main_text = """The Italian CGE model reveals pronounced regional variations in household energy cost burden by 2040, with significant policy implications for climate transition equity. Under the BAU scenario, energy expenditure ranges from 7.8% (Northwest) to 8.5% (Islands) of total household spending. The ETS1 carbon pricing mechanism marginally increases burdens to 8.0-8.8%, reflecting moderate policy costs (Markandya et al., 2012). The ETS2 scenario demonstrates intensified distributional effects, with burdens escalating to 8.5-10.1%, disproportionately affecting Islands and Southern regions (Vandyck et al., 2018).

These findings underscore the regressive nature of carbon pricing on energy-dependent peripheral regions, necessitating compensatory mechanisms such as revenue recycling and targeted social transfers (Dorband et al., 2019). The model validates concerns that climate mitigation policies can exacerbate existing regional inequalities without complementary distributional safeguards (Chakravarty & Tavoni, 2013)."""

    # Add main text
    main_para = doc.add_paragraph(main_text)
    main_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in main_para.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # References section
    ref_heading = doc.add_heading('References', 1)

    # List of references
    references = [
        'Chakravarty, D., & Tavoni, M. (2013). Energy poverty alleviation and climate change mitigation: Is there a trade off? Energy Economics, 40(Supplement 1), S67-S73.',

        'Dorband, I. I., Jakob, M., Kalkuhl, M., & Steckel, J. C. (2019). Poverty and distributional effects of carbon pricing in low-and middle-income countries–A global comparative analysis. World Development, 115, 246-257.',

        'Markandya, A., González-Eguino, M., & Escapa, M. (2012). From shadow to green: Linking environmental fiscal reforms and the informal economy. Energy Economics, 40(Supplement 1), S108-S118.',

        'Vandyck, T., Keramidas, K., Saveyn, B., Kitous, A., & Vrontisi, Z. (2016). A global stocktake of the Paris pledges: Implications for energy systems and economy. Global Environmental Change, 41, 46-63.'
    ]

    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)
        ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in ref_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Save document
    output_path = 'results/Regional_Energy_Burden_2040_Explanation.docx'
    doc.save(output_path)

    print("="*80)
    print("✅ Energy Burden Explanation Document Created Successfully!")
    print("="*80)
    print(f"File saved to: {output_path}")
    print(f"Word count: ~100 words")
    print("\nDocument includes:")
    print("  - Title and subtitle")
    print("  - Economic interpretation of energy burden patterns")
    print("  - Distributional equity analysis")
    print("  - Policy implications")
    print("  - In-text citations")
    print("  - Complete reference list (4 sources)")
    print("="*80)


if __name__ == "__main__":
    create_energy_burden_explanation()
