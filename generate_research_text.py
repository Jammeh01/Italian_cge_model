"""
Generate Research Article Text for CO2 Emission Reductions Figure
================================================================
Creates a professional MS Word document explaining the figure
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


def create_research_article_text():
    """Create Word document with figure explanation"""

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading(
        'Results: CO2 Emission Reduction Pathways under Carbon Pricing Scenarios', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add space
    doc.add_paragraph()

    # Main text - Results section
    heading = doc.add_heading(
        '4.1 Carbon Emission Trajectories and Policy Effectiveness', level=2)

    # Paragraph 1 - Introduction to the figure
    p1 = doc.add_paragraph()
    p1.add_run(
        "Our analysis reveals distinct carbon emission trajectories across the three policy scenarios "
        "examined in this study. Using a recursive-dynamic Computable General Equilibrium (CGE) model "
        "calibrated to Italy's 2021 baseline year (112.9 Mt CO₂), we quantified the emission reduction "
        "potential of progressive carbon pricing policies over the 2021-2040 period (see Figure 1). "
    )
    p1_run = p1.add_run(
        "The model incorporates sector-specific energy intensities, fuel substitution elasticities, and "
        "endogenous technological change through investment-driven renewable capacity expansion "
    )
    p1.add_run("(Lanz & Rausch, 2011; Paltsev et al., 2005).")

    # Paragraph 2 - BAU scenario
    p2 = doc.add_paragraph()
    p2.add_run(
        "Under the Business-As-Usual (BAU) scenario, which assumes continuation of existing policies "
        "without additional carbon pricing measures, our projections indicate a substantial emission "
        "reduction of 40.7% by 2040 relative to the 2021 baseline. This trajectory reflects ongoing "
        "structural economic changes, technological improvements in energy efficiency, and the natural "
        "decarbonization of Italy's energy mix through increased renewable energy penetration "
    )
    p2.add_run("(European Environment Agency, 2023; IEA, 2022). ")
    p2.add_run(
        "The emissions decline from 112.9 Mt CO₂ in 2021 to 80.9 Mt CO₂ by 2030, representing a 28.3% "
        "reduction. However, this reduction trajectory falls significantly short of the European Union's "
        "binding 2030 climate target of at least 55% emission reduction compared to 1990 levels "
    )
    p2.add_run("(European Commission, 2021; Regulation (EU) 2021/1119). ")
    p2.add_run(
        "Our findings suggest that without additional policy intervention, Italy's decarbonization "
        "pathway, while positive, remains insufficient to meet international climate commitments."
    )

    # Paragraph 3 - ETS1 scenario
    p3 = doc.add_paragraph()
    p3.add_run(
        "The implementation of ETS1—a carbon pricing mechanism targeting industrial sectors with an "
        "initial price of €53.90/tCO₂ in 2021, escalating at 4% annually—demonstrates enhanced emission "
        "reduction performance. By 2040, ETS1 achieves a 48.9% reduction from the 2021 baseline, "
        "representing an additional 8.2 percentage points of emission abatement compared to BAU. "
    )
    p3.add_run(
        "This translates to approximately 9.3 Mt CO₂ of avoided emissions in 2040 alone. The carbon "
        "price signal induces fuel switching from coal and gas to electricity, particularly in "
        "energy-intensive industries, while simultaneously incentivizing energy efficiency improvements "
    )
    p3.add_run("(Böhringer & Rutherford, 2008; Rausch et al., 2011). ")
    p3.add_run(
        "Notably, the ETS1 scenario exhibits a steeper emission reduction trajectory after 2025, "
        "suggesting that the cumulative effect of carbon pricing becomes increasingly pronounced as the "
        "price level rises and industries complete capital stock adjustments. By 2030, ETS1 achieves a "
        "29.9% reduction, marginally exceeding the BAU scenario's 28.3% but still falling short of the "
        "EU's 55% target, indicating that sectoral coverage of carbon pricing is a critical determinant "
        "of policy effectiveness "
    )
    p3.add_run("(Tvinnereim & Mehling, 2018).")

    # Paragraph 4 - ETS2 scenario
    p4 = doc.add_paragraph()
    p4.add_run(
        "The most ambitious policy scenario, ETS2, which extends carbon pricing to buildings and "
        "transport sectors starting in 2027 (€45/tCO₂, growing at 2.5% annually), yields substantially "
        "superior emission outcomes. Our model projects a 65.9% reduction by 2040, representing a 25.2 "
        "percentage point improvement over BAU and 17 percentage points over ETS1. This corresponds to "
        "28.4 Mt CO₂ of avoided emissions in 2040 compared to BAU. The accelerated reduction under ETS2 "
        "reflects the high carbon intensity of road transport and residential heating, sectors that "
        "account for approximately 50% of Italy's energy-related emissions "
    )
    p4.add_run("(Eurostat, 2023; ISPRA, 2022). ")
    p4.add_run(
        "The inclusion of these sectors triggers household behavioral responses, including increased "
        "adoption of electric vehicles, heat pump installations, and building retrofitting activities "
    )
    p4.add_run("(Gillingham & Stock, 2018; Gerarden et al., 2017). ")
    p4.add_run(
        "Importantly, our results demonstrate that ETS2 achieves a 35.1% reduction by 2030, the highest "
        "among all scenarios, though still below the 55% EU target. The steeper reduction gradient "
        "post-2027 in the ETS2 scenario (from 35.1% in 2030 to 65.9% in 2040) reflects both the "
        "lagged response of durable goods investments in households and the compound effect of multiple "
        "carbon pricing mechanisms operating simultaneously across the economy."
    )

    # Paragraph 5 - Comparative analysis
    p5 = doc.add_paragraph()
    p5.add_run(
        "A comparative analysis of the three scenarios reveals several critical insights for climate "
        "policy design. First, the marginal emission reduction effectiveness of carbon pricing exhibits "
        "diminishing returns but remains substantial: expanding from no carbon pricing (BAU) to "
        "industrial carbon pricing (ETS1) yields 8.2 percentage points of additional reduction, while "
        "further expanding to buildings and transport (ETS2) delivers an additional 17 percentage points. "
        "Second, all scenarios demonstrate progressive emission reductions over time, indicating that "
        "the Italian economy possesses significant decarbonization potential through both autonomous "
        "technological change and price-induced behavioral adjustments "
    )
    p5.add_run("(Acemoglu et al., 2012; Aghion et al., 2016). ")
    p5.add_run(
        "Third, the temporal dynamics of emission reductions differ markedly: BAU exhibits a relatively "
        "smooth, gradual decline, ETS1 shows moderate acceleration after 2025, while ETS2 displays a "
        "pronounced inflection point at 2027 (policy introduction) followed by rapid reductions. These "
        "patterns align with the heterogeneous capital turnover rates across economic sectors and the "
        "discrete nature of major investment decisions in energy infrastructure "
    )
    p5.add_run("(Kalkuhl et al., 2012).")

    # Paragraph 6 - Policy implications
    p6 = doc.add_paragraph()
    p6.add_run(
        "Our findings carry significant implications for Italy's climate policy architecture and the "
        "broader European Green Deal framework. While the comprehensive ETS2 approach substantially "
        "outperforms more limited interventions, achieving nearly two-thirds emission reduction by 2040, "
        "even this ambitious scenario falls short of the deep decarbonization required for climate "
        "neutrality by 2050 "
    )
    p6.add_run("(IPCC, 2022; European Commission, 2019). ")
    p6.add_run(
        "The analysis suggests that carbon pricing, while necessary, must be complemented by additional "
        "policy instruments including regulatory standards, technology-specific subsidies, and public "
        "infrastructure investments to accelerate the transition "
    )
    p6.add_run("(Stiglitz et al., 2017; Stern, 2007). ")
    p6.add_run(
        "Furthermore, the relatively modest near-term reductions across all scenarios (28-35% by 2030) "
        "underscore the urgency of immediate policy action, as delayed mitigation imposes increasing "
        "costs and reduces the feasible solution space for achieving long-term climate targets "
    )
    p6.add_run("(Rogelj et al., 2018). ")
    p6.add_run(
        "The sectoral coverage of carbon pricing emerges as a crucial design parameter: our results "
        "demonstrate that extending pricing to previously uncovered sectors (buildings and transport) "
        "generates disproportionately large emission reductions, suggesting that comprehensive economy-wide "
        "pricing is essential for cost-effective deep decarbonization."
    )

    # Add figure caption
    doc.add_paragraph()
    caption = doc.add_paragraph()
    caption_run = caption.add_run(
        "Figure 1. CO₂ Emission Reduction Trajectories under Alternative Carbon Pricing Scenarios "
        "(2021-2040). "
    )
    caption_run.bold = True
    caption.add_run(
        "The figure displays percentage emission reductions relative to the 2021 BAU baseline "
        "(112.9 Mt CO₂) for three policy scenarios: Business-As-Usual (BAU, solid blue line), ETS1 "
        "targeting industrial sectors (dashed purple line), and ETS2 extending coverage to buildings "
        "and transport (solid orange line). The horizontal dashed line represents the EU 2030 climate "
        "target of 55% reduction. The vertical dashed line marks 2030. All projections derived from a "
        "recursive-dynamic CGE model calibrated to Italy's 2021 economic structure."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Add space before references
    doc.add_paragraph()
    doc.add_paragraph()

    # References section
    ref_heading = doc.add_heading('References', level=1)

    # Create hanging indent style for references
    references = [
        "Acemoglu, D., Aghion, P., Bursztyn, L., & Hemous, D. (2012). The environment and directed "
        "technical change. American Economic Review, 102(1), 131-166. "
        "https://doi.org/10.1257/aer.102.1.131",

        "Aghion, P., Dechezleprêtre, A., Hemous, D., Martin, R., & Van Reenen, J. (2016). Carbon taxes, "
        "path dependency, and directed technical change: Evidence from the auto industry. Journal of "
        "Political Economy, 124(1), 1-51. https://doi.org/10.1086/684581",

        "Böhringer, C., & Rutherford, T. F. (2008). Combining bottom-up and top-down. Energy Economics, "
        "30(2), 574-596. https://doi.org/10.1016/j.eneco.2007.03.004",

        "European Commission. (2019). The European Green Deal. COM(2019) 640 final. Brussels: European "
        "Commission.",

        "European Commission. (2021). 'Fit for 55': Delivering the EU's 2030 Climate Target on the way "
        "to climate neutrality. COM(2021) 550 final. Brussels: European Commission.",

        "European Environment Agency. (2023). Trends and projections in Europe 2023: Tracking progress "
        "towards Europe's climate and energy targets. EEA Report No 10/2023. Luxembourg: Publications "
        "Office of the European Union.",

        "European Union. (2021). Regulation (EU) 2021/1119 of the European Parliament and of the Council "
        "establishing the framework for achieving climate neutrality (European Climate Law). Official "
        "Journal of the European Union, L 243/1.",

        "Eurostat. (2023). Energy statistics - an overview. Statistics Explained. Luxembourg: European "
        "Commission.",

        "Gerarden, T. D., Newell, R. G., & Stavins, R. N. (2017). Assessing the energy-efficiency gap. "
        "Journal of Economic Literature, 55(4), 1486-1525. https://doi.org/10.1257/jel.20161360",

        "Gillingham, K., & Stock, J. H. (2018). The cost of reducing greenhouse gas emissions. Journal "
        "of Economic Perspectives, 32(4), 53-72. https://doi.org/10.1257/jep.32.4.53",

        "IEA (International Energy Agency). (2022). Italy 2022 Energy Policy Review. Paris: OECD/IEA.",

        "IPCC (Intergovernmental Panel on Climate Change). (2022). Climate Change 2022: Mitigation of "
        "Climate Change. Contribution of Working Group III to the Sixth Assessment Report. Cambridge: "
        "Cambridge University Press. https://doi.org/10.1017/9781009157926",

        "ISPRA (Istituto Superiore per la Protezione e la Ricerca Ambientale). (2022). Italian Greenhouse "
        "Gas Inventory 1990-2020: National Inventory Report 2022. Rome: ISPRA.",

        "Kalkuhl, M., Edenhofer, O., & Lessmann, K. (2012). Learning or lock-in: Optimal technology "
        "policies to support mitigation. Resource and Energy Economics, 34(1), 1-23. "
        "https://doi.org/10.1016/j.reseneeco.2011.08.001",

        "Lanz, B., & Rausch, S. (2011). General equilibrium, electricity generation technologies and the "
        "cost of carbon abatement: A structural sensitivity analysis. Energy Economics, 33(5), 1035-1047. "
        "https://doi.org/10.1016/j.eneco.2011.06.003",

        "Paltsev, S., Reilly, J. M., Jacoby, H. D., Eckaus, R. S., McFarland, J., Sarofim, M., Asadoorian, "
        "M., & Babiker, M. (2005). The MIT Emissions Prediction and Policy Analysis (EPPA) Model: Version 4. "
        "MIT Joint Program on the Science and Policy of Global Change Report 125. Cambridge, MA: MIT.",

        "Rausch, S., Metcalf, G. E., & Reilly, J. M. (2011). Distributional impacts of carbon pricing: "
        "A general equilibrium approach with micro-data for households. Energy Economics, 33(Supplement 1), "
        "S20-S33. https://doi.org/10.1016/j.eneco.2011.07.023",

        "Rogelj, J., Popp, A., Calvin, K. V., Luderer, G., Emmerling, J., Gernaat, D., ... & Tavoni, M. "
        "(2018). Scenarios towards limiting global mean temperature increase below 1.5°C. Nature Climate "
        "Change, 8(4), 325-332. https://doi.org/10.1038/s41558-018-0091-3",

        "Stern, N. (2007). The Economics of Climate Change: The Stern Review. Cambridge: Cambridge "
        "University Press. https://doi.org/10.1017/CBO9780511817434",

        "Stiglitz, J. E., Stern, N., Duan, M., Edenhofer, O., Giraud, G., Heal, G., ... & Winkler, H. "
        "(2017). Report of the High-Level Commission on Carbon Prices. Washington, DC: Carbon Pricing "
        "Leadership Coalition.",

        "Tvinnereim, E., & Mehling, M. (2018). Carbon pricing and deep decarbonisation. Energy Policy, "
        "121, 185-189. https://doi.org/10.1016/j.enpol.2018.06.020"
    ]

    # Add references with hanging indent
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save document
    output_file = 'results/CO2_Emission_Reductions_Figure_Explanation.docx'
    os.makedirs('results', exist_ok=True)
    doc.save(output_file)

    print(f"\n{'='*80}")
    print("RESEARCH ARTICLE TEXT GENERATED")
    print(f"{'='*80}")
    print(f"Document saved to: {output_file}")
    print(f"\nDocument contains:")
    print("  - Professional results section (6 paragraphs)")
    print("  - Figure caption")
    print("  - Complete reference list (21 citations)")
    print("  - Formatted for high-rank journal submission")
    print("  - Total: ~1,500 words")
    print(f"{'='*80}\n")

    return output_file


if __name__ == "__main__":
    create_research_article_text()
